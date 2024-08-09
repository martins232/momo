import json
import pprint
import re
from urllib import response
from django.shortcuts import render, redirect, get_object_or_404
import exam
from users.models import Student, User, Teacher, Grade
from main . decorators import teacher
from teachers. forms import UserUpdateForm,TeacherUpdateForm, ChangeProfilePicture, TopicForm
# from django.contrib.auth.decorators import login_required
from exam.forms import ExamForm, QuestionForm
from django.contrib import messages
from exam .models import Exam, Question, Session
from teachers.models import Subject, Topic
from django.core.exceptions import ObjectDoesNotExist,PermissionDenied
from django.db.models import Q
from django.db.models.signals import post_save, pre_delete
from django.dispatch import receiver

from django.http.response import JsonResponse
from django.db.models import Count, Avg, StdDev
from django.utils.html import strip_tags

from django.http import HttpResponse
from docx import Document
from io import BytesIO
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from django.db import transaction
from django.core.exceptions import ValidationError


from django.utils import timezone
from django.contrib.postgres.aggregates import ArrayAgg



from django.template.loader import render_to_string
from django.views.decorators.cache import cache_page
from openai import OpenAI
from collections import defaultdict
import google.generativeai as genai   
import markdown
genai.configure(api_key="AIzaSyCoc6rtdzgRLlOv1GOpYohbEIORIsrWvhE")


# Create your views here.
# @login_required(login_url="login")
def userProfile(request):
    user = User.objects.get(username = request.user.username)
    subjects = Subject.objects.filter(teacher=user)
    context = {
        "user": user,
        "subjects":subjects,
        "count_course" : subjects.count()
    }
    return render(request, "teachers/profile.html", context)
@teacher
# @login_required(login_url="login")
def editProfile(request, pk):
    user = User.objects.get(id =pk)
    p_form = UserUpdateForm(request, instance = user,)
    t_form = TeacherUpdateForm(request, instance= user.teacher)
    
    
    
    if request.method == "POST":
        p_form = UserUpdateForm(request, request.POST, instance = user)
        t_form = TeacherUpdateForm(request, request.POST, instance= user.teacher)
        if t_form.is_valid():
            p_form.save()
            t_form.save()
            messages.success(request, "Profile updated")
            return redirect("edit-profile", pk =user.id)
        else:
           messages.add_message(request, messages.ERROR, p_form.errors.as_ul())      
           messages.add_message(request, messages.ERROR, t_form.errors.as_ul())      
    context = {
        "p_form": p_form,
        "t_form" : t_form,
        
    }
    return render(request, "teachers/edit_profile.html", context)
@teacher
# @login_required(login_url="login")
def editProfileImage(request, pk):
    
    teacher = Teacher.objects.get(id=pk)
    form = ChangeProfilePicture(instance=teacher)
    if request.method == "POST":
        form = ChangeProfilePicture(request.POST, request.FILES, instance=teacher)
        if form.is_valid():
            form.save()
            return redirect("edit-profile", pk=request.user.id)
    context ={
        "form":form
    }
    return render(request, "teachers/image.html", context)

def dashboard(request):
    now = timezone.now()
    exams = Exam.objects.filter(teacher=request.user, deleted= False)
    grades = Grade.objects.values("grade")
    
    
    context = {
        "total_exam": exams.count(),
        "active": exams.filter(start_date__lt=now, end_date__gt=now).count(),
        "pending": exams.filter(start_date__gt=now).count(),
        "ended": exams.filter(end_date__lt=now).count(),
        "grades": grades
    }
    
    return render(request, "teachers/dashboard.html", context)
@teacher
# @login_required(login_url="login")
def scheduledExam(request):
    #form to create exam
    form = ExamForm(request, request.POST or None)
    
    #get all the exams by this teacher
    # exams = Exam.objects.filter(teacher=request.user, )
    exams = Exam.objects.filter( Q(start_date__gt=timezone.now()) | Q(start_date__lte=timezone.now(), end_date__gt=timezone.now()), teacher=request.user)
    
    #all the classes 
    # grades = Grade.objects.filter(Q(exam__start_date__gt=timezone.now()) | Q(exam__start_date__lte=timezone.now(), exam__end_date__gt=timezone.now()), exam__teacher=request.user).distinct()
    
    unique_grades = exams.values('grade').distinct()
    grades = Grade.objects.filter(id__in=[g['grade'] for g in unique_grades])
    
    
    # # Get the current time once to avoid multiple calls
    # current_time = timezone.now()

    # # Filter exams that are either currently ongoing or will start in the future
    # exams = Exam.objects.filter(
    #     Q(start_date__gt=current_time) | 
    #     Q(start_date__lte=current_time, end_date__gt=current_time), 
    #     teacher=request.user
    # )
    
    # #all the classes 
    # grades = Grade.objects.filter(
    #             exam__in=exams
    #         ).distinct().select_related('exam')
    if request.method == "POST":
        form = ExamForm(request, request.POST, )
        if form.is_valid():
            exam = form.save(commit=False)
            exam.teacher = request.user
            exam.save()
            messages.add_message(request, messages.SUCCESS, "Exam created")
            return redirect("scheduled-exam")
        else:
            messages.add_message(request, messages.ERROR, "Exam not created, check form for error")  
    context = {
        "form": form,
        "exams": exams,
        "grades": grades,
    }
    
    return render(request, "teachers/schedule_exam.html", context)

@teacher
# @login_required(login_url="login")
def closedExam(request):
    """Exams that have been completed"""
    # grades = Grade.objects.filter(exam__end_date__lt=timezone.now(), exam__teacher=request.user).distinct()
    exams = Exam.objects.filter(end_date__lt=timezone.now(), teacher=request.user, deleted=False).order_by("grade").select_related("grade")
    unique_grades = exams.values('grade').distinct()
    grades = Grade.objects.filter(id__in=[g['grade'] for g in unique_grades])
    
    
    
    context = {
        "exams": exams,
        "grades": grades
        }
    return render(request, "teachers/closed_exam.html", context)

@teacher
@receiver(post_save, sender=Subject)
def set_exam_teacher_null(sender, instance, **kwargs):
    """
    This function is used to transfer subject, exams and question from one teacher to another
    """
    # instance is the subject object that was deleted
    # get the exams that are related to this subject
    exams = Exam.objects.filter(subject=instance) #all the exams by the new teacher of this subject
    # # # loop through the exams and set their teacher to null
    for exam in exams:
        exam.teacher = instance.teacher
        exam.save()
        
        # for question in questions:
        #     question.teacher = instance.teacher
        #     question.save()
        #     print(question)
        
@teacher
# @login_required(login_url="login")
def deleteExam(request):
    if is_ajax(request=request):
        id = request.POST.get("id")
        exam = get_object_or_404(Exam, id=id)
        exam.delete()
        return JsonResponse({"message": "deleted"})
        
    else:
        raise PermissionDenied
        
    
@teacher
# @login_required(login_url="login")
def editExam(request, pk):
    exam = get_object_or_404(Exam, id=pk)
    if request.user != exam.teacher:
        return redirect("404")
    form = ExamForm(request, instance=exam, initial={"duration": exam.duration_to_minutes}) 
    
    if request.method == "POST":
        form = ExamForm(request, request.POST, instance=exam) 
        if form.is_valid():
            name= form.cleaned_data["name"]
            form.save()
            messages.add_message(request, messages.SUCCESS, "Exam saved")
            return redirect("scheduled-exam")
        
    context = {
        "form": form,
        "obj_name": "EXAM"
    }  
    return render(request, "teachers/edit.html", context) 



    

@teacher 
# @login_required(login_url="login")
def deleteQuestion(request, pk):
    question = get_object_or_404(Question, id=pk)
    if request.user != question.exam.teacher:
        return redirect("404")
    
    exam = get_object_or_404(Question, id=pk)
    exam.delete()
    messages.add_message(request, messages.SUCCESS, "Question deleted")
    return redirect(request.META.get('HTTP_REFERER', '/'))
    # return redirect("all-questions")
      
    


def sessionDashboardData(request, pk):
    exam = get_object_or_404(Exam, id=pk)
    students = User.objects.filter(is_student=True, student__grade=exam.grade, student__status="Approved").exclude(id__in=Session.objects.filter(exam=exam).values("user"))
    sessions = list(Session.objects.filter(exam=exam).values("user", "score", "elapsed_time", "completed", "misconduct", "time_started"))
    data = []
    for session in sessions:
        session["name"] = User.objects.get(id=session["user"]).get_full_name()
        data.append(session)

    for student in students:
        session = {
            "user": student.id,
            "score": None,
            "elapsed_time": None,
            "completed": None,
            "misconduct": None,
            "time_started": None,
            "name": student.get_full_name(),
        }
        data.append(session)   
    return JsonResponse({"pass_mark":exam.pass_mark, "rows": data})
    
####################################################################
def sessionDashboard(request, pk):
    exam = Exam.objects.get(id=pk)
    
    questions = Question.objects.filter(exam=exam)
    total_student = exam.grade.student_set.filter(status="Approved").count()
    misconduct = Session.objects.filter(exam = exam, misconduct=True).count()
    completed = Session.objects.filter(exam = exam, completed=True).count()
    passed = Session.objects.filter(exam = exam,score__gte=exam.pass_mark, completed=True).count()
    
    
    try:
        avg_score = Session.objects.filter(exam=exam, completed=True).aggregate(score=Avg("score"))["score"]
        avg_score = round(avg_score, 1)
    except Exception as e:
        avg_score = "-"
    # dict(Session.objects.filter(exam=1).values_list("misconduct").annotate(count=Count("misconduct", distinct=True)))
    
    context ={
        "exam": exam,
        "exam_id": exam.id,
        "avg_score": avg_score,
        "questions": questions,
        "total_student": total_student,
        "misconduct": misconduct,
        "completed": completed,
        "passed": passed,
        "fail" : completed - passed,
        
    }
    return render(request, "teachers/session_dashboard.html", context)

def studentPerformance(request, pk):
    student_name = request.GET.get("student")
    id = request.GET.get("id")
    
    exam = Exam.objects.get(id=pk)
    student = User.objects.get(id=id)
    session = Session.objects.get(exam=exam, user=student)
    choices = session.choices
    
    # print(choices)
    # assert()
    
    questions = exam.question_set.all().values()
            
    data_ =[]
    
    
    no_correct = 0
    no_incorrect = 0
    no_unanswered = 0
    

    for question in questions:
        options = [ question["option_A"], question["option_B"], question["option_C"], question["option_D"] ]
        # zipper =dict(list( zip(['A', 'B', 'C', 'D'], options)))
        # answer_Abbrv =  question["answer"] # the answer in the database e.g "C"
        
        if choices[str(question["id"])][0] == "":
            no_unanswered += 1
        if choices[str(question["id"])][0] == question["answer"]:
            no_correct += 1
        if choices[str(question["id"])][0] != question["answer"]:
            no_incorrect += 1
        
        
        data_.append({"question": question["question"], 
                      "options_A" : question["option_A"], 
                      "options_B" : question["option_B"], 
                      "options_C" : question["option_C"], 
                      "options_D" : question["option_D"], 
                      "answer": question["answer"], 
                      "choice": choices[str(question["id"])][0]})
    
    
    context = {
        "student": student,
        "exam": exam,
        "student_name": student_name,
        "session" : session,
        "data_": data_,
        "no_unanswered": no_unanswered,
        "no_correct": no_correct,
        "no_incorrect": no_incorrect
  
    }
    return render(request, "teachers/student_performance.html", context)

        
#JSON data after exams has been taken
def examDashboardData(request, pk):
    exam = Exam.objects.get(id = pk)
    sessions = Session.objects.filter(exam= exam).select_related('user')
    questions = exam.get_all_question.order_by("id")
    topics_count, questions_without_topic = exam.topic_count()
   
    data = []
    for question in questions:
        correct_count = 0
        incorrect_count = 0
        unanswered_count = 0
        
        correct_student = {}
        incorrect_student = {}
        unanswered_student = {}
        
        options = {"A" :question.option_A, "B" :question.option_B, "C" :question.option_C, "D" :question.option_D,} #{"A":"........"}
        #dict to get number of student that chose an option
        option_choice_nos={question.option_A:0, question.option_B:0, question.option_C:0, question.option_D:0} #{"......":0}
        
        for session in sessions:
            student_choice = session.choices[str(question.id)][0]
            if student_choice == "":
                unanswered_count += 1
                unanswered_student[str(session.user.id)] = str(session.user.get_full_name())
                
            else:
                option_text = options[student_choice]
                option_choice_nos[option_text]+=1
                if student_choice == question.answer:
                    correct_count +=1
                    correct_student[str(session.user.id)] = str(session.user.get_full_name())
                    

                if student_choice != question.answer:
                    incorrect_count +=1
                    incorrect_student[str(session.user.id)] = session.user.get_full_name()
                    
        if question.topics:
            topics_count[question.topics.name]["correct"] += correct_count
            topics_count[question.topics.name]["incorrect"] += incorrect_count
            topics_count[question.topics.name]["unanswered"] += unanswered_count
            topics_count[question.topics.name]["no_questions"] += 1
                
        data.append({
            "id":question.id,
            "question": question.question, 
            "answer": options[question.answer],
            "correct": correct_count, 
            "incorrect": incorrect_count, 
            "unanswered": unanswered_count,
            "correct_student": correct_student,
            "incorrect_student": incorrect_student,
            "unanswered_student": unanswered_student,
            "option_count" : option_choice_nos
            })
      
    return JsonResponse({"rows":data, "topics": topics_count})

def newExamDashboard(request, pk):
    exam = Exam.objects.get(id = pk)
    sessions = Session.objects.filter(exam= exam).order_by("-score")
    context ={
        "exam":exam,
        "sessions":sessions,
    }
    return render(request, "teachers/new_exam_dashboard.html", context)
  
    
#dashboard after exams has been taken  
def examDashboard(request, pk):
    exam = Exam.objects.get(id = pk)
    sessions = Session.objects.filter(exam= exam).order_by("-score")
    
    
    context ={
        "exam":exam,
        "sessions":sessions,
        "pk": pk,
        "session_count": sessions.count()
    }
    
       
    return render(request, "teachers/exam_dashboard.html", context)


def questionData(request):
    subject = request.GET.get("subject") # used in the drag page to filter questions for a particular subject 
    search = request.GET.get("search") 
    limit = int(request.GET.get("limit"))
    offset = int(request.GET.get("offset"))
    unassigned = request.GET.get("unassigned")
    exam_id = request.GET.get("exam")
    # unassigned = bool(int(request.GET.get("unassigned", 0)))
    
    # if there is a particular subject that is making the request through ajax
    if subject != None:
        subject=int(subject)
        length = Question.objects.filter(subject=subject).count() #get all the questions by the logged in teacher
        questions = Question.objects.filter(subject=subject).filter(Q(subject__name__icontains=search) | Q(topics__name__icontains=search) | Q(question__icontains=search))
        
       
        """Note unassigned here are questions that do to belong 'to a particular' exam """
        unassigned = bool(int(unassigned))
        if unassigned: #this block finds question not attached to this exam
            questions = questions.exclude(exam = exam_id).values("id", "question","topics__id", "topics__name","option_A", "option_B","option_C","option_D", "answer","subject", "subject__name","explanation")
            length =  questions.count()
        else:
            questions = questions.filter(exam = exam_id).values("id", "question","topics__id", "topics__name","option_A", "option_B","option_C","option_D", "answer","subject", "subject__name", "exam__id", "explanation")
            length =  questions.count()
        
    else:
         # for the sake of other subjects that would use this JSON data
        subject_list = Subject.objects.filter(teacher=request.user) # used in the all_question page to filter all questions by a particular teacher 
        length = Question.objects.filter(subject__in=subject_list).count() #get all the questions by the logged in teacher
        questions = Question.objects.filter(subject__in=subject_list).filter(Q(subject__name__icontains=search) | Q(topics__name__icontains=search) | Q(question__icontains=search)).order_by("-created").annotate(exam__name = ArrayAgg("exam__name", distinct=True)).values("id", "question","topics__id", "topics__name","option_A", "option_B","option_C","option_D", "answer","subject","explanation", "subject__name", "exam__name")
    
    
        
        
        
    

    
    
    
    
    
    if search: # when searching fields using the search input
        length =  questions.count()
        
    
    
   
       
    questions = questions[offset:limit+offset]
    return JsonResponse({"total":length, "offset": offset, "rows":list(questions)})


def is_ajax(request): #check if a call is an ajax call
    return request.META.get('HTTP_X_REQUESTED_WITH') == 'XMLHttpRequest'

def createQuestion(request):  
    if is_ajax(request=request):
        form = QuestionForm(request, request.POST)
        if form.is_valid():
            form.save()
            return JsonResponse({"message":"added"})
        else:
            errors = form.errors 
            return JsonResponse({"message": "Validation errors", "errors": errors})
            
        
def deleteQuestion(request):
    if is_ajax(request=request):
        ids = json.loads(request.POST.get("id"))
        quest = Question.objects.filter(id__in = ids)
        a = quest.delete()
        return JsonResponse({"deleted": True})
    
def editQuestion(request):
    if is_ajax(request=request):
        id = request.POST.get("id")
        question = Question.objects.get(id=id)
        form = QuestionForm(request, request.POST, instance=question)
        if form.is_valid():
            form.save()
            return JsonResponse({"res":"added"})
        else:
            errors = form.errors 
            return JsonResponse({"message": "Validation errors", "errors": errors})
        
        
        
def viewExam(request, pk):
    exam = Exam.objects.get(id=pk)
    
    context = {"exam": exam} #used for sending subject id in the templated
    if  exam.get_exam_status == "active" and  exam.ready is False:
        context["show_no_question_warning"] = True
        
    if exam.get_exam_status =="ended" or "active":
        grades = Grade.objects.all()
    
        # Create a dictionary to hold topics grouped by grade
        topics_grouped_by_grade = {}

        for grade in grades:
            # Fetch topics for each grade
            topics = Topic.objects.filter(grade=grade, subject=exam.subject)
            if topics.count()>0:
                topics_grouped_by_grade[grade] = topics
        
        context["topics_grouped_by_grade"] = topics_grouped_by_grade
        
    if exam.get_exam_status =="pending" or "ended":
        context["update_topic"] = True
    return render(request, "teachers/view_exam.html", context)

def assignQuestionToExam(request, pk):
    exam = Exam.objects.get(id=pk)
    if is_ajax(request=request):
        ids = json.loads(request.POST.get("id"))
        questions = Question.objects.filter(id__in = ids)
        for question in questions:
            question.exam.add(exam)
        
        if exam.get_no_question > 0:           
            exam.ready = True
            exam.save()
        return JsonResponse({"message": f"{len(ids)} Question(s) added to {exam.name}. <br>Total questions: <span class='fw-bold'>{exam.get_no_question}</span>"})
    
def removeQuestionFromExam(request, pk):
    exam = Exam.objects.get(id=pk)
    if is_ajax(request=request):
        ids = json.loads(request.POST.get("id"))
        questions = Question.objects.filter(id__in = ids)
        for question in questions:
            question.exam.remove(exam)
        
        if exam.get_no_question < 1:
            exam.ready = False
            exam.save()
        
        
        return JsonResponse({"message": f"{len(ids)} Question(s) removed from {exam.name}. <br>Total questions: <span class='fw-bold'>{exam.get_no_question}</span>"})

def topicsData(request):
    teacher = request.user
    subject = Subject.objects.filter(teacher = teacher)
    data = {}
    grade = request.GET.get("grade") 
    
    
    if grade is not None:
        for sub in subject:
            data[sub.name] = list(Topic.objects.filter(subject=sub, grade=grade).values("grade","id", "subject", "name"))
    return JsonResponse(data)

def allTopics(request):
    form = TopicForm(request)
    grades = Grade.objects.all()
    # subjects = Subject.objects.filter(teacher = request.user)
    context = {
        "form": form,
        "grades": grades
        }
    
    return render(request, "teachers/topics.html", context)

def addTopic(request):
    if is_ajax(request=request):
        form = TopicForm(request, request.POST)
        if form.is_valid():
            topic = form.save()
            msg = topic.subject.name + " for " + str(topic.grade)
            return JsonResponse({"message": f"New topic created in {msg}"})
        else:
            return JsonResponse({"message":"No"})
        
def editTopic(request):
    if is_ajax(request=request):      
        topic_id = request.POST.get("id")
        topic = get_object_or_404(Topic, id=topic_id)       
        form = TopicForm(request, request.POST, instance= topic)
        if form.is_valid():
            topic =form.save()
            msg = topic.subject.name + " for " + str(topic.grade)
            return JsonResponse({"message": f"Topic updated in {msg}"})
        else:
            return JsonResponse({"message": form.errors})


def deleteTopic(request):
    if is_ajax(request=request):
        id = request.POST.get("id")
        topic =Topic.objects.get(pk=id)
        if topic:
            msg = topic.subject.name + " for " + str(topic.grade)
            topic.delete()
            return JsonResponse({"message":f"Topic deleted from {msg}"})
        else:
            return JsonResponse({"message":"Something went wrong"})
        

#-------------------------------------------------------------
@teacher
# @login_required(login_url="login")
def allQuestions(request):
    
    
    form = QuestionForm(request)
    context = {
        "form":form,
        # "subjects": subjects, 
    }
    
    return render(request, "teachers/all_questions.html", context)

# @cache_page(60*15)
def question_create(request):
    data = dict()
    form = QuestionForm(request, request.POST or None)
    if request.method == "POST":
        if form.is_valid():
            form.save()
            data['form_is_valid'] = True
        else:
            data['form_is_valid'] = False
            data['html_form'] = form
    context = {'form': form}
    data['html_form'] = render_to_string('teachers/includes/create_question.html',context,request=request)
    return JsonResponse(data)

    

# # @cache_page(60*15)
def question_edit(request, pk):
    question = Question.objects.get(id=pk)
    data = dict()
    form = form = QuestionForm(request, request.POST or None, instance=question)
    if request.method == "POST":
        if form.is_valid():
            form.save()
            data['form_is_valid'] = True
        else:
            data['form_is_valid'] = False
            data['html_form'] = form
    context = {'form': form}
    data['html_form'] = render_to_string('teachers/includes/create_question.html',context,request=request)
    return JsonResponse(data)

def question_delete(request):
    """Function to delete questions through ajax"""
    ids = request.POST.get("ids")
    ids = json.loads(ids)
    questions = Question.objects.filter(id__in =ids)
    questions.delete()
    return JsonResponse({"response":True})

def update_topics(request):
    """Function to delete questions through ajax"""
    ids = request.POST.get("ids")
    topic = request.POST.get("topic")
    ids = json.loads(ids)
    topic = json.loads(topic)
    questions = Question.objects.filter(id__in =ids)
    questions.update(topics=topic)                           
    return JsonResponse({"response":True})

def convertToDocx(request):
    #to convert questions already in the database to docx file
    document = Document()
    """Function to delete questions through ajax"""
    ids = request.POST.get("ids")
    ids = json.loads(ids)
    question_instances = Question.objects.filter(id__in =ids)
    # question_instances = Question.objects.filter(subject__teacher=request.user)
    questions = [
        (
            strip_tags(q.question),
            [f"a) {q.option_A   }", f"b) {q.option_B}", f"c) {q.option_C}", f"d) {q.option_D}"]
        )
        for q in question_instances
    ]  
    i = 1   
    for question, choices in questions:
        para = document.add_paragraph()

        # Add the question text
        para.add_run(str(i) + "\t" +question+ "\n").bold = True

        # Add choices
        for choice in choices:
            para.add_run("\t" + choice +"\n")  # Using a tab character to align choices
            
        i+=1    
    
    buffer = BytesIO() #creates a virtual memory in RAM for reading and writing byte
    document.save(buffer) #write the document’s bytes into the in-memory buffer instead of a physical file.
    buffer.seek(0) #it “seeks” to the 0th byte (the start) of the buffer, so you can read from it from the beginning.
    
    # Set up the HTTP response
    response = HttpResponse(
        buffer.getvalue(), #gets the store data from the buffer virtual memory in bytes
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document' #This tells the client’s browser that the server is sending a Word document.
    )
    response['Content-Disposition'] = 'attachment; filename="emeka.docx"'    
    return response





# Function to add a areas where for question set to the document
def add_question_set(doc, question_prefix):
    # Add question with indentation
    # Define the question and options
    question_text = ""

    options = ["A)  ", "B)  ", "C)  ", "D)  "]
    
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT # type: ignore
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(1))  # type: ignore # Add a tab stop at 1.5 cm
    
    run = p.add_run(question_prefix)
    run.font.size = Pt(11) # type: ignore
    run.add_tab()  # Insert tab character
    run = p.add_run(question_text)
    run.font.size = Pt(11) # type: ignore

    set_paragraph_border(p) # function to add border to each paragraph
    
    # Add table with options
    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    
    # Adjust column widths
    table.columns[0].width = Inches(2.5) # type: ignore
    table.columns[1].width = Inches(2.5) # type: ignore

    # Add options to the table
    for i, option in enumerate(options):
        cell = table.cell(i // 2, i % 2)
        cell.text = option
        cell.paragraphs[0].runs[0].font.size = Pt(11) # type: ignore

    last_row = table.rows[-1]
    merged_cell = last_row.cells[0].merge(last_row.cells[1])
    merged_cell.text = "Ans)   "
    merged_cell.paragraphs[0].runs[0].font.size = Pt(11)
    
    # Set table borders
    set_table_borders(table)
    
    # Set table indentation
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.left_indent = Cm(1)  # type: ignore # Indent by 1 cm

    # Add some spacing after each question set
    doc.add_paragraph("\n")
    
def set_paragraph_border(paragraph):
    # Get the paragraph properties element
    p = paragraph._p
    pPr = p.get_or_add_pPr()

    # Create the border element
    pbdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 1/8 point, 8 = 1 point
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        pbdr.append(border)

    # Append the border element to paragraph properties
    pPr.append(pbdr)
    
def set_table_borders(table):
    tbl = table._element
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 1/8 point, 8 = 1 point
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)

    tbl.tblPr.append(tblBorders)
    
def prepareDocxTemplateForExam(request):
    #when the teacher submits the form form for template creation
    if request.method =="POST":
        no_of_questions = int(request.POST.get("nos"))
        title = request.POST.get("title")
        
        # Create a new Document
        doc = Document()
        
        # Add level 2 heading
        title = doc.add_paragraph("{} question template".format(title.title()))
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set the color to red
            run.font.size = Pt(20)

        # Add level 3 heading
        level3_heading = doc.add_paragraph("Warning: Do not alter the structure of this template")
        level3_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in level3_heading.runs:
            run.font.color.rgb = RGBColor(255, 0, 0)  # Set the color to red
        
        # Add three question sets to the document
        for i in range(1,no_of_questions +1):
            add_question_set(doc, f"{i}.",)

        # Save the document
        buffer = BytesIO()
        doc.save(buffer) #write the document’s bytes into the in-memory buffer instead of a physical file.
        buffer.seek(0) #it “seeks” to the 0th byte (the start) of the buffer, so you can read from it from the beginning.
        # Set up the HTTP response
        response = HttpResponse(
            buffer.getvalue(), #gets the store data from the buffer virtual memory in bytes
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document' #This tells the client’s browser that the server is sending a Word document.
        )
        response['Content-Disposition'] = 'attachment; filename="question_upload.docx"'    
        return response


def bulk_create_questions(subject_id, exam_ids, data):
    '''After all validations has been made bulk create a template docx file'''
    questions = []
    errors = []
    for item in data:
        question_instance = Question(
            subject=subject_id,
            question=item[0],
            option_A=item[1][0],
            option_B=item[1][1],
            option_C=item[1][2],
            option_D=item[1][3],
            answer=item[1][4].upper()
        )
        try:
            # Validate each instance
            question_instance.full_clean()
            questions.append(question_instance)
        except ValidationError as e:
            # Collect validation errors
            errors.append((question_instance, e))

    
    if errors:
        return errors    

    
    # Use transaction to ensure atomicity
    with transaction.atomic():
        created_questions = Question.objects.bulk_create(questions)
        # Add ManyToMany relationships
        for question in created_questions:
            question.exam.add(exam_ids)

        
def handleUploadedDocx(request):
    from docx.oxml.table import CT_Tbl
    
    if request.method =="POST" and 'file' in request.FILES:
        uploaded_file = request.FILES['file']
        exam_name = request.POST.get("title")
        exam = Exam.objects.get(name=exam_name)
        subject = exam.subject
        
        if exam.get_exam_status == "active" and exam.ready:
            return JsonResponse({"message": False, "reason": "<b>Upload error</b>:<br> Exam is already active. Adding question will disrupt exam"})
        
        if not uploaded_file.name.endswith('.docx'):
            return JsonResponse({"message": False, "reason": "<b>File error</b>:<br> File type is not supported. Please upload a .docx file. You can get one by downloading a template"})
        # Use python-docx to process the uploaded .docx file
        doc = Document(uploaded_file)
        
        question_pattern = re.compile(r'^\d+\.\s')
    
        # Regular expression to match the option tags
        # option_pattern = re.compile(r'^[A-Z]\)\s*')
        option_pattern = re.compile(r'^(?:[aA-zZ]\)|Ans\))\s*')
        questions = []
        options = []
        
        
        current_question = ''
        collecting = False
        
        
        for paragraph in doc.paragraphs:
            # Check if the paragraph text matches the question pattern
            if question_pattern.match(paragraph.text):
                if current_question:
                    questions.append(current_question.strip())
                current_question = question_pattern.sub('', paragraph.text).strip()
                collecting = True
            elif collecting:
                if not paragraph.text.strip():  # Skip empty lines
                    continue
                current_question += '<br>' + paragraph.text.strip() 
            else:
                continue

        if current_question:
            questions.append(current_question.strip())
        
        
        
        # Extract options from tables
        for table in doc.tables:
            table_options = []
            for row in table.rows:
                for cell in row.cells:
                    # Collect cell text without duplicates and strip option tags
                    option_text = option_pattern.sub('', cell.text.strip())
                    if option_text and option_text not in table_options:
                        table_options.append(option_text.strip())
            if table_options:
                options.append(table_options)        
        
        
        
        # check that 
        if len(questions)>0:
            if len (questions) == len(options):
                message = True
                questions_and_options = list(zip(questions, options))
                                
                for question_no , questions_and_option in enumerate(questions_and_options):
                    if len(questions_and_option[1]) != 5:
                        message = False #this will notify teachers of incomplete option in one or more options
                        break
                # print(questions_and_options)
                errors = bulk_create_questions(subject, exam, questions_and_options) 
                if errors:
                    # Prepare error messages for the user
                    error_messages = []
                    for instance, error in errors:
                        error_messages.append(f"Question: '<b>{instance.question}</b>' has errors: {error.messages}")
                    
                    return JsonResponse({"message": False, "errors": error_messages})   
                if message:
                    exam.ready = True
                    exam.save()
                    return JsonResponse({"message": True}) 
                else:
                    return JsonResponse({"message": False, "reason": "<b>Incomplete Question Option</b>:<br> One or more options for a question are not fully completed. Please review the file and provide all necessary information. Error occurred at <b>question {}</b>".format(question_no + 1)})
            else:
                return JsonResponse({"message": False, "reason": "<b>Question Length Mismatch</b>:<br> Please ensure you are using the template file correctly. If you believe the template has been modified incorrectly, please download a fresh copy."}) 
        else:
            return JsonResponse({"message": False, "reason": "<b>No question found</b>:<br> We were unable to find any questions in the uploaded file. This might be due to the file's content not being recognized or the file being empty. Please check the file and try again."})     
            
        
        
        #     fullText.append(para.text)
        # document = '\n'.join(fullText)
        # print(document)
              
    
    
def changeStudentsPassword(request):
    '''
    Function that allows teachers to reset students password upon request
    '''
    students = Student.objects.filter(request_password=True)
    if request.method == "POST":
        ids = request.POST.get("ids")
        ids = json.loads(ids)
        students = Student.objects.filter(id__in = ids)
        for student in students:
            student.user.set_password("12345678")
            student.user.save()
        students.update(request_password=False)
        return JsonResponse({"msg": True})
    context = {
        "students": students
    }
    return render(request, "teachers/change_students_password.html", context)


def explanationWithAI(request):
    if request.method == "POST":
        id = request.POST.get("id")
        question = Question.objects.get(id=id)
        
        if question:
            system_message = """
                    You are an educational assistant. When provided with a multiple-choice question and the correct answer, generate a detailed explanation that helps students understand the concept behind the question and why the correct answer is accurate. The explanation should be clear, concise, and educational using simple terms. Start the response with the correct answer.

                    If you determine that the provided correct answer is incorrect, notify the teacher by stating, "The provided answer is incorrect," and then proceed to give the correct answer along with the explanation.
                    
                    """
            
            model = genai.GenerativeModel('gemini-1.5-pro-latest', system_instruction=system_message)
            
            data = f"""
                {question.question}\n
                A) {question.option_A}\n
                B) {question.option_B}\n
                C) {question.option_C}\n
                D) {question.option_D}\n
                Ans) {question.answer}"""
            response = model.generate_content(data)
            html = markdown.markdown(response.text)

        return JsonResponse({"status": True, "response":html})
    # return explanation
    
    


def SavegenerateExplanation(request):
    
    if request.method == "POST":
        id = request.POST.get("id")
        explanation = request.POST.get("explanation")
        question = Question.objects.get(id=id)
        question.explanation = explanation
        question.save()
        
        return JsonResponse({"message":True})

def generateExamSummary(request):
 
    model = genai.GenerativeModel('gemini-1.5-pro-latest', 
            system_instruction=
                    """
                    
                    **System Instruction:**

                    You are an expert data analyst. Your task is to analyze the exam data provided as a Python dictionary. Your response should be clear, elaborate, and easily understood by a teacher. Use the provided example report as a guide, and feel free to add any important analysis you think may be necessary. Ensure your response is helpful for the teacher, using clear and simple language. Provide explanations where necessary and avoid technical jargon. Note that you are not to expose the way the data is organized to the teacher.


                    **Example Report: Chemistry (CS 101) Exam Report - For Grade 1**
                    **For : [name of Teacher]
                    1. **Overall Performance Summary**
                    - Provide an overview of the class's performance, including average scores, highest and lowest scores, average time taken to complete the exam. and any other relevant statistics and identify any notable observations if there is. Determine factors influencing the pass rate by examining potential causes.

                    2. **Score Distribution**
                    - Analyze the score distribution and also check if there is correlation between score and number of attempt. Describe how the scores are spread out across the group, identifying notable patterns or observations. Calculate and interpret key descriptive statistics including mean, median, mode, standard deviation, and range. Identify and analyze outliers. Compare the distribution to relevant benchmarks or standards. Explore potential relationships between the score distribution and other variables. Determine the shape of the distribution (normal, skewed, bimodal, etc.) and calculate percentiles to understand performance levels. Provide insights for the teacher into the implications of the score distribution for overall performance and suggest potential actions to address identified disparities. 
                        Here is an example:
                        The distribution of scores in your class is quite spread out, with a standard deviation of 30.11. This indicates a significant variation in student performance.

                        * Mean (Average): [value] - Represents the average score achieved by the class.
                        * Median: [value] - Represents the middle score when all scores are arranged in order.
                        * Mode: [value] - The most frequent score achieved by students.
                        * Range: [lowest score - highest score]- Highlights the difference between the lowest and highest scores.
                        The score distribution appears to be moderately skewed, with a few low scores potentially pulling the average down.

                        * Implications and Actions:

                            - Address the needs of students who are struggling. Consider providing additional support or individualized instruction to help them catch up.
                            - Challenge high-achieving students. Offer them extension activities or opportunities to deepen their understanding.

                    3. **Difficulty Distribution**
                    - Analyze question performance data to determine question difficulty and overall assessment reliability. Using the questions difficulty indices for each question that i have already calculated for you, Note that the higher the indices the easier the question. classify these questions into difficulty levels (easy, medium, difficult) based on established thresholds of your discretion. Note that higher difficulty indices indicate easier questions. Identify questions that deviate significantly from the target difficulty level. Analyze the distribution of question difficulties to assess the overall balance of the assessment. Provide recommendations for improving question quality and overall assessment reliability.
                    
                        Here is an example for an exam:
                        Based on the provided difficulty indices:

                        - Easy Questions (Index above 80%): Questions 208, 211, 216, 242, and 246 seem to be easy as a high percentage of students answered them correctly.
                        - Medium Difficulty Questions (Index between 50% - 80%): A large portion of the questions fell within this range, including questions like 207, 212, 214, 217, 219, 220, and others.
                        - Difficult Questions (Index below 50%): Questions 225, 234, 235, 237, and 241 appear to be more challenging, as indicated by their lower difficulty indices.
                        The exam seems to have a good balance of difficulty levels, which is important for a comprehensive assessment. However, you could consider reviewing the difficult questions to identify if there were any specific concepts causing trouble for many students or if there were any issues with the clarity of those questions.

                    4. **Topic Performance**
                    - Evaluate how students performed on different topics covered in the exam. Identify areas where students excelled and areas needing improvement and provide teachers on foundational concepts students must know to excel in these topics that student didn't do well.
                    
                        Here is an example for a chemistry exam:
                        A preliminary analysis suggests potential variations in student performance across different topics. While a detailed breakdown requires further information on individual student scores for each question, the following observations can be made:

                        * Topics with potential areas of weakness include "Environmental Chemistry", "Chemical Kinetics", "Redox Reactions", and      "States of Matter". These topics had a lower proportion of correct answers compared to others.
                        * Foundational concepts to focus on within these challenging topics include:
                        - Environmental Chemistry: Understanding the impact of chemicals on the environment and how chemical reactions affect pollution.
                        - Chemical Kinetics: Grasping the factors influencing reaction rates, such as temperature and catalysts.
                        - Redox Reactions: Identifying oxidation and reduction processes, and balancing redox equations.
                        - States of Matter: Understanding the properties of different states of matter and the factors influencing phase changes.
                        it was observed that some topics where categorized as "No topic". Do well to categorize these questions in the topic analysis section. Providing students with additional support and resources on these specific topics may be beneficial.

                    5. **Correlation Between Time Spent and Performance**
                    - Investigate if there is a relationship between the time students spent on the exam and their performance. 
                    
                    Here is an example when there is no correlation:
                    A preliminary analysis reveals that time spent on the exam does not guarantee high scores. Notably, some students with shorter completion times achieved high scores, while others who spent more time still struggled.
                    This discrepancy hints at the influence of additional factors, such as mastery of the subject matter, effective test-taking approaches, guessing accuracy or anxiety levels, which can outweigh the impact of time spent on the exam.

                    6. **Potential Instances of Misconduct**
                    - Look for any irregularities or patterns that might suggest cheating or other forms of misconduct.

                    7. **Recommendations**
                    - Write recommendation for the teacher based on your exam report. Identify areas of strength and weakness, and suggest specific strategies for improving student performance. Consider both short-term and long-term recommendations, and provide actionable advice for enhancing instruction, addressing knowledge gaps, and promoting student success.

                    
                    
                
    """)  
                    # 8. **Overall Summary and Next Steps**
                    # - Conclude with a summary of the key findings and propose next steps for the teacher.
    if request.method == "POST":
        exam = Exam.objects.get(id = request.POST.get("id"))
        sessions = Session.objects.filter(exam=exam)
         # Get all questions for the given exam
        questions = Question.objects.filter(exam=exam)
        
        # Create a dictionary to store questions by topic
        questions_by_topic = defaultdict(list)
        
        # Iterate through the questions and organize them by topic
        for question in questions:
            topic_name = question.topics.name if question.topics else "No Topic"
            questions_by_topic[topic_name].append(question.id)
        
        average_score = sessions.aggregate(Avg('score'))['score__avg']
        average_elapsed_time = sessions.aggregate(Avg('elapsed_time'))['elapsed_time__avg']
        std_dev_score = sessions.aggregate(StdDev('score'))['score__stddev']
        pass_rate = sessions.filter(score__gte=exam.pass_mark).count() / sessions.count() * 100
         # Collecting data for score distribution
        scores = list(sessions.values('attempts', 'score'))
        # Collecting detailed choices data
        choices_data = [session.choices for session in sessions]
        
        # Analyzing performance by question
        question_performance = {}
        for choices in choices_data:
            for question_id, answer_data in choices.items():
                if question_id not in question_performance:
                    question_performance[question_id] = {'correct': 0, "incorrect": 0, "unanswered":0, 'total': 0}
                if answer_data[1]['status'] == 'correct':
                    question_performance[question_id]['correct'] += 1
                elif answer_data[1]['status'] == 'incorrect':
                    question_performance[question_id]['incorrect'] += 1
                elif answer_data[1]['status'] == 'unanswered':
                    question_performance[question_id]['unanswered'] += 1
                question_performance[question_id]['total'] += 1
        
        # Calculate difficulty for each question
        question_difficulty = {
            question_id: round(data['correct'] / (data["correct"] + data["incorrect"]), 2) * 100
            for question_id, data in question_performance.items()
        }
        # Time vs. Performance and elapsed time analysis
        time_vs_performance = list(sessions.values('elapsed_time', 'score'))
        
        misconduct_count = sessions.filter(misconduct=True).count()
        data = {
            "exam_name":exam.name,
            "exam_class(grade)":exam.grade.grade,
            "subject": exam.subject.name,
            "teacher": exam.teacher.first_name,
            'overall_performance': {
                "pass_mark": exam.pass_mark,
                'average score of the whole class': round(average_score,2),
                'std dev score': std_dev_score,
                'pass rate': pass_rate,
                "exam duration": exam.seconds_to_hms,
                "average_elapsed_time_in_seconds(you are to convert this to the format of exam duration)":round(average_elapsed_time,2),
                'scores': scores, 
            },
            "question performance": question_performance,
            'question difficulty indices(%)': question_difficulty,
            "topics and numbers of question it has": dict(questions_by_topic),
            'time of each student spent vs performance': time_vs_performance,
            'elapsed times of students in seconds': [tp['elapsed_time'] for tp in time_vs_performance],
            'misconduct_count': misconduct_count,
        }
        
        # pprint.pp(data)
        
        response = model.generate_content(f"{data}")
        html = markdown.markdown(response.text)
        
    return JsonResponse({"summary": html})